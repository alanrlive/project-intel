"""
Create sample test files for batch upload tests.

Run from repo root:
    python tests/create_test_data.py

Generates:
    tests/test_data/raid_example.xlsx   — RAID log spreadsheet
    tests/test_data/meeting_notes.docx  — meeting notes with actions
    tests/test_data/budget_overview.pdf — minimal PDF with financial content
"""

import struct
import zlib
from pathlib import Path

TEST_DATA_DIR = Path(__file__).parent / "test_data"
TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)


# ── XLSX ──────────────────────────────────────────────────────────────────────

def create_raid_xlsx():
    try:
        import openpyxl
    except ImportError:
        print("[ERROR] openpyxl not installed. Run: pip install openpyxl")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "RAID Log"

    headers = ["Type", "Description", "Owner", "Impact", "Status"]
    ws.append(headers)

    rows = [
        ["Risk",   "Vendor delivery delayed by 2+ weeks due to supply chain issues",
         "Alan R",  "High",   "Open"],
        ["Risk",   "Budget approval may slip past April 15 deadline",
         "Finance", "Medium", "Open"],
        ["Action", "Review and sign vendor proposal",
         "Alan R",  "High",   "Open"],
        ["Action", "Schedule kickoff meeting with stakeholders",
         "PM",      "Medium", "Open"],
        ["Issue",  "Legacy API endpoint returning 504 errors intermittently",
         "Dev Lead","High",   "In Progress"],
        ["Dependency", "Budget sign-off must happen before vendor contract",
         "Finance", "High",   "Blocking"],
    ]
    for row in rows:
        ws.append(row)

    dest = TEST_DATA_DIR / "raid_example.xlsx"
    wb.save(dest)
    print(f"[OK] Created {dest}")


# ── DOCX ──────────────────────────────────────────────────────────────────────

def create_meeting_notes_docx():
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] python-docx not installed. Run: pip install python-docx")
        return

    doc = Document()
    doc.add_heading("Project Kickoff Meeting — 5 April 2026", 0)

    doc.add_heading("Attendees", 1)
    doc.add_paragraph("Alan R (PM), Sarah K (Dev Lead), Mike T (Finance), Lisa P (QA)")

    doc.add_heading("Action Items", 1)
    actions = [
        "Alan to finalise vendor contract by April 12 (HIGH priority)",
        "Sarah to set up CI/CD pipeline by April 20 (HIGH priority)",
        "Mike to approve Q2 budget allocation by April 10",
        "Lisa to write test plan for phase 1 by April 18 (MEDIUM priority)",
    ]
    for a in actions:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Deadlines", 1)
    deadlines = [
        "Phase 1 development complete: May 1 2026",
        "User acceptance testing: May 15 2026",
        "Go-live: June 1 2026",
    ]
    for d in deadlines:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("Risks Identified", 1)
    doc.add_paragraph(
        "Resource availability risk: Dev team split across two projects until April 30. "
        "Medium impact, medium likelihood. Mitigation: prioritise this project for the first two weeks."
    )
    doc.add_paragraph(
        "Scope creep risk: Stakeholders requesting additional reporting features not in original plan. "
        "High impact, high likelihood. Mitigation: strict change control process."
    )

    doc.add_heading("Dependencies", 1)
    doc.add_paragraph("Budget approval enables vendor contract signing.")
    doc.add_paragraph("Vendor contract blocks project kickoff.")

    dest = TEST_DATA_DIR / "meeting_notes.docx"
    doc.save(dest)
    print(f"[OK] Created {dest}")


# ── PDF (minimal, pure-Python) ────────────────────────────────────────────────

def _deflate(data: bytes) -> bytes:
    return zlib.compress(data)[2:-4]  # strip zlib header/checksum — raw deflate


def create_budget_pdf():
    """
    Build a minimal but valid PDF from scratch using only stdlib.
    Contains enough financial text for the LLM to extract from.
    """
    text = (
        "Project Budget Overview - Q2 2026\n\n"
        "Total approved budget: GBP 450,000\n\n"
        "Breakdown:\n"
        "  - Vendor licences:       GBP 120,000  (due April 30 2026)\n"
        "  - Development resources: GBP 200,000\n"
        "  - Infrastructure:        GBP  80,000\n"
        "  - Contingency (10%):     GBP  50,000\n\n"
        "Risks:\n"
        "  1. Vendor costs may increase by 15% if contract not signed by April 12.\n"
        "     Impact: High. Likelihood: Medium. Owner: Finance team.\n"
        "  2. Cloud infrastructure spend may exceed estimate by 20% due to traffic spike.\n"
        "     Impact: Medium. Likelihood: Low. Mitigation: Set billing alerts at 80% threshold.\n\n"
        "Actions:\n"
        "  - Finance to raise PO for vendor by April 10 (HIGH priority, due April 10 2026)\n"
        "  - DevOps to configure cost alerts by April 15 (MEDIUM priority)\n"
        "  - PM to get CFO sign-off on contingency release by April 20\n\n"
        "Deadlines:\n"
        "  - Q2 budget freeze: April 30 2026\n"
        "  - Mid-year review: July 1 2026\n"
    )

    # Encode text as a PDF content stream (just BT/ET blocks per line)
    lines = text.split("\n")
    stream_parts = ["BT", "/F1 11 Tf", "50 750 Td", "14 TL"]
    for line in lines:
        safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream_parts.append(f"({safe}) Tj T*")
    stream_parts.append("ET")
    stream_content = "\n".join(stream_parts).encode("latin-1")

    # Build PDF objects as plain bytes (no compression for simplicity)
    objects: list[bytes] = []

    # obj 1: catalog
    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    # obj 2: pages
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    # obj 3: page
    objects.append(
        b"3 0 obj\n"
        b"<< /Type /Page /Parent 2 0 R\n"
        b"   /MediaBox [0 0 612 792]\n"
        b"   /Contents 4 0 R\n"
        b"   /Resources << /Font << /F1 5 0 R >> >> >>\n"
        b"endobj\n"
    )
    # obj 4: content stream
    stream_len = len(stream_content)
    objects.append(
        f"4 0 obj\n<< /Length {stream_len} >>\nstream\n".encode()
        + stream_content
        + b"\nendstream\nendobj\n"
    )
    # obj 5: font
    objects.append(
        b"5 0 obj\n"
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"
        b"endobj\n"
    )

    # Build cross-reference table
    header = b"%PDF-1.4\n"
    body = b""
    offsets: list[int] = []
    pos = len(header)
    for obj in objects:
        offsets.append(pos)
        body += obj
        pos += len(obj)

    xref_pos = len(header) + len(body)
    xref = f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"

    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    )

    pdf_bytes = header + body + xref.encode() + trailer.encode()

    dest = TEST_DATA_DIR / "budget_overview.pdf"
    dest.write_bytes(pdf_bytes)
    print(f"[OK] Created {dest}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Writing test files to: {TEST_DATA_DIR}\n")
    create_raid_xlsx()
    create_meeting_notes_docx()
    create_budget_pdf()
    print("\n[OK] Done. Run: python tests/test_batch_upload.py")
